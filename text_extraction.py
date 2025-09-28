import fitz  # PyMuPDF
from docx import Document
import streamlit as st
import os

# Extract text from PDF using PyMuPDF
def extract_text_from_pdf(file_path):
    full_text = ""
    try:
        with fitz.open(file_path) as pdf_doc:
            for page_num in range(len(pdf_doc)):
                page = pdf_doc.load_page(page_num)
                blocks = page.get_text("blocks")
                for block in sorted(blocks, key=lambda b: (b[1], b[0])):
                    full_text += block[4] + "\n"
        return full_text.strip()
    except Exception as e:
        return f"❌ Error reading PDF: {e}"

# Extract text from DOCX (paragraphs + tables)
def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            full_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)

        return "\n".join(full_text)
    except Exception as e:
        return f"❌ Error reading DOCX: {e}"

# Main Streamlit app
def main():
    st.set_page_config(page_title="Text Extractor", layout="centered")
    st.title("📄 Document Text Extractor")
    st.markdown("Upload a `.pdf` or `.docx` file to extract **all text content**, including from tables.")

    uploaded_file = st.file_uploader("📤 Upload File", type=["pdf", "docx"])

    if uploaded_file is not None:
        file_name = uploaded_file.name
        file_ext = os.path.splitext(file_name)[1].lower()
        temp_path = f"temp_uploaded_file{file_ext}"

        # Save uploaded file
        with open(temp_path, "wb") as f:
            f.write(uploaded_file.read())

        with st.spinner(f"⏳ Extracting text from `{file_name}`..."):
            if file_ext == ".pdf":
                extracted_text = extract_text_from_pdf(temp_path)
            elif file_ext == ".docx":
                extracted_text = extract_text_from_docx(temp_path)
            else:
                st.error("Unsupported file type.")
                return

        # Clean up
        os.remove(temp_path)

        st.success("✅ Text extraction complete!")
        st.subheader("📝 Extracted Text")
        st.text_area("Full Text Output", extracted_text, height=400)

        st.subheader("📊 Text Statistics")
        st.write(f"**File Name:** {file_name}")
        st.write(f"**Characters Extracted:** {len(extracted_text)}")
        st.write(f"**Words Extracted:** {len(extracted_text.split())}")

        st.download_button(
            label="💾 Download Extracted Text",
            data=extracted_text,
            file_name=f"extracted_{os.path.splitext(file_name)[0]}.txt",
            mime="text/plain"
        )

if __name__ == "__main__":
    main()
